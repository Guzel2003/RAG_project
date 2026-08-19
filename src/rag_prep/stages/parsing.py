"""Парсинг документов через Unstructured + python-docx для .docx."""

from __future__ import annotations

import csv
import logging
from typing import Any

from unstructured.partition.auto import partition

from ..config import ParserConfig
from ..models import ParseFailure, ParseResult, RawElement, SourceFile
from ..utils import stable_id

LOGGER = logging.getLogger(__name__)

SECTION_TYPES = {"Title", "Header"}
TITLE_TYPES = {"Title"}


class UnstructuredParsingStage:
    """Парсит PDF/TXT/HTML через Unstructured, CSV вручную, DOCX через python-docx."""

    def __init__(self, config: ParserConfig, default_section: str):
        self.config = config
        self.default_section = default_section

    def run(self, sources: list[SourceFile]) -> ParseResult:
        elements: list[RawElement] = []
        failures: list[ParseFailure] = []

        for source in sources:
            try:
                elements.extend(self._parse_source(source))
            except Exception as exc:
                if self.config.fail_on_error:
                    raise
                LOGGER.exception("Failed to parse %s", source.source)
                failures.append(ParseFailure(
                    source=source.source,
                    file_name=source.file_name,
                    file_type=source.file_type,
                    error_type=exc.__class__.__name__,
                    error_message=str(exc),
                ))

        LOGGER.info("Parsed %d elements from %d files; %d failures",
                    len(elements), len(sources), len(failures))
        return ParseResult(elements=elements, failures=failures)

    def _parse_source(self, source: SourceFile) -> list[RawElement]:
        if source.file_type == "csv":
            return self._parse_csv(source)
        if source.file_type == "docx":
            return self._parse_docx(source)
        if source.file_type == "txt":
            return self._parse_txt_with_fallback(source)
        if source.file_type == "pdf":
            return self._parse_pdf_pymupdf(source)

        # Unstructured для HTML и других форматов
        parsed = partition(
            filename=str(source.path),
            strategy=self.config.strategy,
            encoding=self.config.encoding,
            languages=self.config.languages,
            pdf_infer_table_structure=self.config.pdf_infer_table_structure,
            skip_infer_table_types=self.config.skip_infer_table_types,
        )

        section_path = [self.default_section]
        raw_elements: list[RawElement] = []

        for idx, element in enumerate(parsed):
            text = str(element).strip()
            if not text:
                continue

            elem_type = getattr(element, "category", element.__class__.__name__)
            if elem_type in SECTION_TYPES:
                section_path = self._next_section_path(section_path, text, elem_type)

            section = section_path[-1] if section_path else self.default_section
            raw_elements.append(RawElement(
                source_file=source,
                element_id=self._element_id(source, idx),
                element_index=idx,
                text=text,
                element_type=elem_type,
                section=section,
                section_path=list(section_path),
                metadata=self._metadata_to_dict(getattr(element, "metadata", None)),
            ))

        return raw_elements

    def _parse_pdf_pymupdf(self, source: SourceFile) -> list[RawElement]:
        """Парсит PDF через PyMuPDF (без тяжёлых зависимостей)."""
        import fitz

        doc = fitz.open(str(source.path))
        raw_elements: list[RawElement] = []

        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text("text").strip()

            if not text:
                continue

            raw_elements.append(RawElement(
                source_file=source,
                element_id=self._element_id(source, page_num),
                element_index=page_num,
                text=text,
                element_type="PDFPage",
                section=self.default_section,
                section_path=[self.default_section],
                metadata={"page_number": page_num + 1, "page_count": len(doc)},
            ))

        doc.close()
        LOGGER.info("Parsed PDF %s: %d pages", source.file_name, len(raw_elements))
        return raw_elements

    def _parse_txt_with_fallback(self, source: SourceFile) -> list[RawElement]:
        """Парсит TXT с автоматическим определением кодировки."""
        encodings_to_try = ["utf-8", "cp1251", "koi8-r", "iso-8859-5"]

        text = None
        used_encoding = None

        for enc in encodings_to_try:
            try:
                text = source.path.read_text(encoding=enc)
                used_encoding = enc
                break
            except UnicodeDecodeError:
                continue

        if text is None:
            # Fallback: читаем с игнорированием ошибок
            text = source.path.read_text(encoding="utf-8", errors="ignore")
            used_encoding = "utf-8 (with errors ignored)"

        LOGGER.info("Parsed %s with encoding: %s", source.file_name, used_encoding)

        # Разбиваем на элементы по абзацам
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]

        raw_elements: list[RawElement] = []
        for idx, para in enumerate(paragraphs):
            raw_elements.append(RawElement(
                source_file=source,
                element_id=self._element_id(source, idx),
                element_index=idx,
                text=para,
                element_type="Paragraph",
                section=self.default_section,
                section_path=[self.default_section],
                metadata={"encoding": used_encoding},
            ))

        return raw_elements

    def _parse_docx(self, source: SourceFile) -> list[RawElement]:
        """Парсит .docx через python-docx."""
        from docx import Document

        doc = Document(str(source.path))
        elements: list[RawElement] = []
        idx = 0

        # Абзацы
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            elements.append(RawElement(
                source_file=source,
                element_id=self._element_id(source, idx),
                element_index=idx,
                text=text,
                element_type="Paragraph",
                section=self.default_section,
                section_path=[self.default_section],
                metadata={},
            ))
            idx += 1

        # Таблицы
        if self.config.docx_include_tables:
            sep = self.config.docx_table_separator
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [" ".join(c.text.strip().split()) for c in row.cells]
                    if any(cells):
                        rows.append(sep.join(cells))
                if rows:
                    text = "Таблица:\n" + "\n".join(rows)
                    elements.append(RawElement(
                        source_file=source,
                        element_id=self._element_id(source, idx),
                        element_index=idx,
                        text=text,
                        element_type="Table",
                        section=self.default_section,
                        section_path=[self.default_section],
                        metadata={"table_rows": len(rows)},
                    ))
                    idx += 1

        return elements

    def _parse_csv(self, source: SourceFile) -> list[RawElement]:
        with source.path.open("r", encoding=self.config.encoding, newline="") as f:
            sample = f.read(65536)
            f.seek(0)
            try:
                dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
            except csv.Error:
                dialect = csv.excel

            reader = csv.DictReader(f, dialect=dialect)
            columns = reader.fieldnames or []
            elements: list[RawElement] = []

            for row_num, row in enumerate(reader, start=1):
                normalized = {str(k).strip(): "" if v is None else str(v).strip()
                              for k, v in row.items() if k is not None}
                text = "\n".join(f"{k}: {v}" for k, v in normalized.items() if v).strip()
                if not text:
                    continue

                elements.append(RawElement(
                    source_file=source,
                    element_id=self._element_id(source, row_num - 1),
                    element_index=row_num - 1,
                    text=text,
                    element_type="CSVRow",
                    section=self.default_section,
                    section_path=[self.default_section],
                    metadata={"csv_row_number": row_num, "csv_columns": columns},
                ))

        return elements

    def _next_section_path(self, current: list[str], title: str, elem_type: str) -> list[str]:
        if elem_type in TITLE_TYPES:
            return [title]
        if current:
            return [current[0], title]
        return [title]

    @staticmethod
    def _element_id(source: SourceFile, index: int) -> str:
        return stable_id(source.source_hash, index)

    @staticmethod
    def _metadata_to_dict(metadata: Any) -> dict[str, Any]:
        if metadata is None:
            return {}
        if hasattr(metadata, "to_dict"):
            return metadata.to_dict()
        if isinstance(metadata, dict):
            return metadata
        return {}